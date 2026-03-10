from __future__ import annotations

import json
import logging
import re
from typing import List

from pydantic import BaseModel, Field
from app.infrastructure.llm import get_llm_gateway

logger = logging.getLogger(__name__)
SYSTEM_PROMPT = """
You are ExamExtractor-1, a high-precision parser for university exam papers.

Your task is to convert a question paper into structured JSON.

-------------------------
IMPORTANT RULES
-------------------------

1. Extract ALL questions from the document.

2. Preserve the exact wording of each question.

3. Ignore instructions, page numbers, and formatting artifacts.

4. Question numbering must match the paper.

Examples:
1
1a
1(a)
2
3
4(b)

Use:
questionNumber = main number
questionSubPart = letter only if it is a SHORT independent question.

-------------------------
SECTION A TYPE
-------------------------

Example:

1 a
1 b
1 c
...
1 j

Each subpart is an independent question.

Output:

questionNumber: 1
questionSubPart: "a"
maxMarks: 2

-------------------------
LONG QUESTIONS
-------------------------

Example:

2 a
2 b

These are NOT separate questions.

They belong to the SAME question.

Output:

questionNumber: 2
questionSubPart: null

questionText must contain BOTH:

(a) text
(b) text

maxMarks = total marks for the question.

-------------------------
OR QUESTIONS
-------------------------

Example:

2 a
2 b

OR

3 a
3 b

This is ONE question choice.

Output ONE question:

questionNumber: 2

questionText:

(a) ...
(b) ...

OR

(a) ...
(b) ...

maxMarks = marks for ONE choice only.

Never output Q2 and Q3 separately.

-------------------------
MARKS
-------------------------

Extract marks if shown:

[5M]
(10 marks)
10M

If two parts each have 5M → maxMarks = 10.

-------------------------
OUTPUT FORMAT
-------------------------

{
 "title": "string",
 "subject": "string",
 "questions": [
   {
     "questionNumber": number,
     "questionSubPart": "optional",
     "questionText": "full text",
     "maxMarks": number,
     "rubric":[
       {
         "description":"string",
         "maxMarks":number
       }
     ]
   }
 ]
}

Return JSON only.
"""


# ------------------------------------------------
# DATA MODELS
# ------------------------------------------------

class RubricItem(BaseModel):
    description: str
    max_marks: float = Field(alias="maxMarks")
    model_config = {"populate_by_name": True}


class ExtractedQuestion(BaseModel):
    question_number: int | None = Field(default=None, alias="questionNumber")
    question_number_or: int | None = Field(default=None, alias="questionNumberOr")
    question_sub_part: str | None = Field(default=None, alias="questionSubPart")
    question_text: str = Field(alias="questionText")
    max_marks: float = Field(alias="maxMarks")
    rubric: List[RubricItem] = Field(default_factory=list)
    rubric_second_option: List[RubricItem] | None = Field(default=None, alias="rubricSecondOption")

    model_config = {"populate_by_name": True}


class ExtractedExam(BaseModel):
    title: str = "Untitled Exam"
    subject: str = "General"
    questions: List[ExtractedQuestion] = Field(default_factory=list)

    model_config = {"populate_by_name": True}


# ------------------------------------------------
# RUBRIC NORMALIZATION
# ------------------------------------------------

def normalize_rubrics(data: dict) -> dict:
    """
    Normalize rubric keys returned by LLM.
    """

    for q in data.get("questions", []):
        if not isinstance(q, dict):
            continue
        rubric = q.get("rubric", [])
        if not isinstance(rubric, list):
            q["rubric"] = []
            continue
        normalized = []

        for r in rubric:

            if not isinstance(r, dict):
                continue

            description = (
                r.get("description")
                or r.get("criteria")
                or r.get("criterion")
            )

            marks = (
                r.get("maxMarks")
                or r.get("marks")
                or r.get("score")
            )

            if description is not None:
                # Remove duplicate "(X marks)" so we don't show "(5 marks) (5 marks)"
                desc_str = str(description).strip()
                desc_str = re.sub(
                    r"(\s*\(\s*\d+(?:\.\d+)?\s*marks?\s*\))\s*\1+",
                    r"\1",
                    desc_str,
                    flags=re.IGNORECASE,
                )
                normalized.append({
                    "description": desc_str,
                    "maxMarks": marks if marks is not None else 0,
                })

        q["rubric"] = normalized

    return data


# ------------------------------------------------
# OR QUESTION MERGING (GENERIC)
# ------------------------------------------------

def merge_or_questions(questions: List[ExtractedQuestion]) -> List[ExtractedQuestion]:

    merged = []
    i = 0

    while i < len(questions):

        q = questions[i]

        if i + 1 < len(questions):

            nxt = questions[i + 1]

            if (
                q.question_number
                and nxt.question_number
                and nxt.question_number == q.question_number + 1
                and q.max_marks == nxt.max_marks
                and q.max_marks >= 5
                and q.question_sub_part is None
                and nxt.question_sub_part is None
            ):

                combined_text = (
                    f"{q.question_text}\n\n"
                    f"OR\n\n"
                    f"(Question {nxt.question_number})\n"
                    f"{nxt.question_text}"
                )

                merged.append(
                    ExtractedQuestion(
                        questionNumber=q.question_number,
                        questionNumberOr=nxt.question_number,
                        questionText=combined_text,
                        maxMarks=q.max_marks,
                        rubric=q.rubric or [],
                        rubricSecondOption=nxt.rubric or []
                    )
                )

                i += 2
                continue

        merged.append(q)
        i += 1

    return merged




# ------------------------------------------------
# MAIN EXTRACTION
# ------------------------------------------------

def extract_exam_from_text(
    question_paper_text: str,
    rubric_text: str | None = None,
    stated_max_marks: int | None = None,
    *,
    merge: bool = True,
) -> ExtractedExam:

    gateway = get_llm_gateway()

    prompt = f"""
QUESTION PAPER TEXT

{question_paper_text}
"""

    if rubric_text:
        prompt += f"\n\nRUBRIC TEXT\n{rubric_text}"

    response = gateway.complete(
        system_prompt=SYSTEM_PROMPT,
        user_prompt=prompt,
        temperature=0,
        max_tokens=8000,
    )

    content = response.content

    # handle block responses
    if isinstance(content, list):

        content = "\n".join(
            block.get("text", "")
            if isinstance(block, dict)
            else str(block)
            for block in content
        )

    content = content.strip()

    if content.startswith("```"):
        content = "\n".join(content.split("\n")[1:-1])

    try:

        data = json.loads(content)

        data = normalize_rubrics(data)

        exam = ExtractedExam.model_validate(data)

        if merge:
            exam.questions = merge_or_questions(exam.questions)

        return exam

    except Exception as e:

        logger.error("Exam extraction failed: %s", e)
        raise


# ------------------------------------------------
# DOCX EXTRACTION
# ------------------------------------------------

def extract_text_from_docx(file_bytes: bytes) -> str:

    from io import BytesIO
    from docx import Document

    doc = Document(BytesIO(file_bytes))

    lines = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)


# ------------------------------------------------
# FAST PDF TEXT EXTRACTION
# ------------------------------------------------

_PDF_FAST_EXTRACT_MIN_CHARS = 80


def extract_text_from_pdf_fast(pdf_path: str) -> str:

    try:
        from pypdf import PdfReader

        reader = PdfReader(pdf_path)

        parts = []

        for i, page in enumerate(reader.pages):

            text = page.extract_text()

            if text and text.strip():
                parts.append(f"--- Page {i+1} ---\n{text.strip()}")

        return "\n\n".join(parts)

    except Exception as e:

        logger.warning("Fast PDF extraction failed: %s", e)
        return ""


def extract_text_from_pdf_via_vision(pdf_path: str) -> str:
    """
    Extract text from PDF: try fast path (pypdf) first.
    For scanned/handwritten PDFs, use OpenAI Vision with one API call per page.
    """
    fast_text = extract_text_from_pdf_fast(pdf_path)
    if len(fast_text.strip()) >= _PDF_FAST_EXTRACT_MIN_CHARS:
        logger.info("Using fast PDF text extraction (no Vision API calls)")
        return fast_text

    import os
    import shutil
    import tempfile
    from pdf2image import convert_from_path

    images = convert_from_path(pdf_path, dpi=150)
    if not images:
        return ""

    logger.info("Using Vision API: %d page(s), one API call per page", len(images))
    gateway = get_llm_gateway()
    tmpdir = tempfile.mkdtemp()
    all_text = []
    try:
        for i, img in enumerate(images, start=1):
            page_path = os.path.join(tmpdir, f"page_{i}.png")
            img.save(page_path, "PNG")
            response = gateway.vision_extract_text(page_path, detail="high")
            all_text.append(f"--- Page {i} ---\n{response.content.strip()}")
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)

    return "\n\n".join(all_text)


# ------------------------------------------------
# IMAGE OCR
# ------------------------------------------------

def extract_text_from_image_via_vision(image_path: str) -> str:

    gateway = get_llm_gateway()

    response = gateway.vision_extract_text(image_path)

    content = response.content

    if isinstance(content, str):
        return content.strip()

    if isinstance(content, list):

        return "\n".join(
            block.get("text", "")
            if isinstance(block, dict)
            else str(block)
            for block in content
        ).strip()

    return str(content).strip()